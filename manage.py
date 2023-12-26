#!/usr/bin/env python
"""Django's command-line utility for administrative tasks."""
import os
import shutil
import sys
import pandas as pd
import psycopg2
import openpyxl
from docx import Document
import yadisk
import requests
import schedule


def main():
    # connparams = {
    #     'dbname': 'postgres',
    #     'user': 'postgres',
    #     'password': 'root1324',
    #     'host': 'localhost',
    #     'port': '5432'
    # }
    # #
    # conn = psycopg2.connect(**connparams)
    # tablename = 'base_task'
    # query = f'SELECT * FROM {tablename};'
    # df = pd.read_sql(query, conn)
    #
    # doc = Document()
    # doc.add_heading('Данные из таблицы', level=1)
    # for index, row in df.iterrows():
    #     doc.add_paragraph(str(row))
    # doc.save('base_task.docx')
    #
    # df.to_json('base_task.json', orient='records')
    #
    # source_db_path = 'base_task.json'
    # backup_db_path = 'base_task_backup.json'
    # shutil.copyfile(source_db_path, backup_db_path)
    #
    # disk = yadisk.YaDisk(token='y0_AgAAAABVAAZHAAsKAAAAAAD10ZGIFnrOA5ZVQiiSzxaUlhlm9GiLIQM')
    # disk.remove("base_task_backup.json", permanently=True)
    # disk.upload("base_task_backup.json", f"/base_task_backup.json")
    #
    # tablename = 'auth_user'
    # query = f'SELECT * FROM {tablename};'
    # df = pd.read_sql(query, conn)
    #
    #
    # doc = Document()
    # doc.add_heading('Данные из таблицы', level=1)
    # for index, row in df.iterrows():
    #     doc.add_paragraph(str(row))
    # doc.save('auth_user.docx')
    #
    # df.to_json('auth_user.json', orient='records')
    #
    # source_db_path = 'auth_user.json'
    # backup_db_path = 'auth_user_backup.json'
    # shutil.copyfile(source_db_path, backup_db_path)
    #
    # disk.upload("auth_user_backup.json", f"/auth_user_backup.json")
    #
    # tablename = 'base_grouptasks'
    # query = f'SELECT * FROM {tablename};'
    # df = pd.read_sql(query, conn)
    #
    # doc = Document()
    # doc.add_heading('Данные из таблицы', level=1)
    # for index, row in df.iterrows():
    #     doc.add_paragraph(str(row))
    # doc.save('base_grouptasks.docx')
    #
    # df.to_json('base_grouptasks.json', orient='records')
    #
    # source_db_path = 'base_grouptasks.json'
    # backup_db_path = 'base_grouptasks_backup.json'
    # shutil.copyfile(source_db_path, backup_db_path)
    #
    # # schedule.every().day.at("00:00").do(disk.upload("output_backup.json", f"/output_backup.json"))
    #
    # disk.upload("base_grouptasks_backup.json", f"/base_grouptasks_backup.json")
    # conn.close()
    #
    # print("success")

    """Run administrative tasks."""
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'todo_list.settings')
    try:
        from django.core.management import execute_from_command_line
    except ImportError as exc:
        raise ImportError(
            "Couldn't import Django. Are you sure it's installed and "
            "available on your PYTHONPATH environment variable? Did you "
            "forget to activate a virtual environment?"
        ) from exc
    execute_from_command_line(sys.argv)


if __name__ == '__main__':
    main()
